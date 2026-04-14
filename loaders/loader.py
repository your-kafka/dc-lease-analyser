from pathlib import Path


def load_pdf(path: str) -> str:
    import pdfplumber

    pages = []
    has_text = False

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            text = page.extract_text()

            if text and text.strip():
                has_text = True
                pages.append(f"[Page {i}]\n{text.strip()}")

    # 🔥 AFTER loop finishes
    if not has_text:
        print("⚠️ No text found, using OCR...")
        return extract_text_from_scanned_pdf(path)

    return "\n\n".join(pages)


def load_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


def load_text(path: str) -> str:
    for encoding in ["utf-8", "utf-8-sig", "latin-1", "cp1252"]:
        try:
            with open(path, "r", encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    raise ValueError(f"Could not decode file: {path}")


LOADERS = {
    ".pdf":  load_pdf,
    ".docx": load_docx,
    ".doc":  load_docx,
    ".txt":  load_text,
    ".md":   load_text,
}


def load_document(path: str) -> tuple[str, str]:
    """
    Detects file type and returns (text, file_type).
    Raises ValueError for unsupported or empty files.
    """
    ext = Path(path).suffix.lower()

    if ext not in LOADERS:
        raise ValueError(f"Unsupported file type '{ext}'. Supported: {list(LOADERS)}")

    text = LOADERS[ext](path)

    if not text or not text.strip():
        raise ValueError(f"No text extracted from '{path}'. File may be a scanned image.")

    return text, ext.lstrip(".")

def extract_text_from_scanned_pdf(path):
    from pdf2image import convert_from_path
    import pytesseract

    images = convert_from_path(path)
    text = ""

    for i, img in enumerate(images):
        text += f"\n[Page {i+1}]\n"
        text += pytesseract.image_to_string(img)

    return text